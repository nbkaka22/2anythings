"""PDF转DOCX转换器插件

将PDF文件转换为Word文档格式
"""

import os
import fitz  # PyMuPDF
from docx import Document
from pdf2docx import parse
import tempfile
import logging
from typing import Dict, Any, List

# 导入基类
from converters.converter_interface import ConverterInterface, ConverterMetadata

logger = logging.getLogger('pdf_converter')

class PDFToDocxConverter(ConverterInterface):
    """PDF转DOCX转换器
    
    使用PyMuPDF和pdf2docx库将PDF文件转换为Word文档
    """
    
    def __init__(self):
        self._temp_files = []
    
    @property
    def name(self) -> str:
        return "pdf_to_docx"
    
    @property
    def description(self) -> str:
        return "将PDF文件转换为Word文档格式，保留文本和基本格式"
    
    @property
    def supported_input_formats(self) -> List[str]:
        return ["pdf"]
    
    @property
    def supported_output_formats(self) -> List[str]:
        return ["docx"]
    
    @property
    def version(self) -> str:
        return "1.0.0"
    
    def validate_input(self, input_path: str) -> bool:
        """验证PDF文件是否有效"""
        try:
            if not os.path.exists(input_path):
                logger.error(f"文件不存在: {input_path}")
                return False
            
            if not input_path.lower().endswith('.pdf'):
                logger.error(f"文件格式不正确，需要PDF文件: {input_path}")
                return False
            
            # 尝试打开PDF文件
            doc = fitz.open(input_path)
            if doc.page_count == 0:
                logger.error(f"PDF文件为空: {input_path}")
                doc.close()
                return False
            
            doc.close()
            return True
            
        except Exception as e:
            logger.error(f"PDF文件验证失败: {e}")
            return False
    
    def convert(self, input_path: str, output_path: str, **kwargs) -> bool:
        """执行PDF到DOCX的转换
        
        Args:
            input_path: PDF文件路径
            output_path: 输出DOCX文件路径
            **kwargs: 转换选项
                - method: 转换方法 ('pdf2docx' 或 'pymupdf')
                - pages: 要转换的页面范围，格式如 '1-3,5,7-9'
                - start_page: 起始页码（从0开始）
                - end_page: 结束页码
        
        Returns:
            bool: 转换是否成功
        """
        try:
            # 获取转换选项
            method = kwargs.get('method', 'pdf2docx')
            pages = kwargs.get('pages', None)
            start_page = kwargs.get('start_page', 0)
            end_page = kwargs.get('end_page', None)
            
            logger.info(f"开始PDF转DOCX转换: {input_path} -> {output_path}")
            logger.info(f"转换方法: {method}")
            
            # 确保输出目录存在
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            if method == 'pdf2docx':
                return self._convert_with_pdf2docx(input_path, output_path, start_page, end_page)
            elif method == 'pymupdf':
                return self._convert_with_pymupdf(input_path, output_path, start_page, end_page)
            else:
                logger.error(f"不支持的转换方法: {method}")
                return False
                
        except Exception as e:
            logger.error(f"PDF转DOCX失败: {e}")
            return False
    
    def _convert_with_pdf2docx(self, input_path: str, output_path: str, start_page: int = 0, end_page: int = None) -> bool:
        """使用pdf2docx库进行转换"""
        try:
            logger.info("使用pdf2docx方法转换")
            
            # 获取PDF页数
            doc = fitz.open(input_path)
            total_pages = doc.page_count
            doc.close()
            
            if end_page is None:
                end_page = total_pages - 1
            
            # 执行转换
            parse(input_path, output_path, start=start_page, end=end_page)
            
            # 验证输出文件
            if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                logger.info(f"PDF转DOCX成功: {output_path}")
                return True
            else:
                logger.error("转换完成但输出文件无效")
                return False
                
        except Exception as e:
            logger.error(f"pdf2docx转换失败: {e}")
            return False
    
    def _convert_with_pymupdf(self, input_path: str, output_path: str, start_page: int = 0, end_page: int = None) -> bool:
        """使用PyMuPDF进行转换（提取文本）"""
        try:
            logger.info("使用PyMuPDF方法转换")
            
            # 打开PDF文档
            pdf_doc = fitz.open(input_path)
            
            if end_page is None:
                end_page = pdf_doc.page_count - 1
            
            # 创建Word文档
            word_doc = Document()
            
            # 逐页提取文本
            for page_num in range(start_page, min(end_page + 1, pdf_doc.page_count)):
                page = pdf_doc[page_num]
                text = page.get_text()
                
                if text.strip():
                    # 添加页面标题
                    if page_num > start_page:
                        word_doc.add_page_break()
                    
                    # 按段落分割文本
                    paragraphs = text.split('\n\n')
                    for paragraph in paragraphs:
                        if paragraph.strip():
                            word_doc.add_paragraph(paragraph.strip())
                else:
                    # 如果没有文本，添加提示
                    word_doc.add_paragraph(f"[第{page_num + 1}页：未检测到文本内容]")
            
            pdf_doc.close()
            
            # 保存Word文档
            word_doc.save(output_path)
            
            # 验证输出文件
            if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                logger.info(f"PDF转DOCX成功: {output_path}")
                return True
            else:
                logger.error("转换完成但输出文件无效")
                return False
                
        except Exception as e:
            logger.error(f"PyMuPDF转换失败: {e}")
            return False
    
    def get_default_options(self) -> Dict[str, Any]:
        """获取默认转换选项"""
        return {
            'method': 'pdf2docx',  # 默认使用pdf2docx方法
            'start_page': 0,
            'end_page': None,
            'preserve_formatting': True,
            'extract_images': False
        }
    
    def cleanup(self):
        """清理临时文件"""
        for temp_file in self._temp_files:
            try:
                if os.path.exists(temp_file):
                    os.remove(temp_file)
            except Exception as e:
                logger.warning(f"清理临时文件失败: {temp_file}, {e}")
        
        self._temp_files.clear()
    
    def _check_pdf_integrity(self, pdf_path: str) -> bool:
        """检查PDF文件完整性"""
        try:
            doc = fitz.open(pdf_path)
            # 尝试访问第一页
            if doc.page_count > 0:
                page = doc[0]
                # 尝试获取页面内容
                page.get_text()
            doc.close()
            return True
        except Exception as e:
            logger.warning(f"PDF完整性检查失败: {e}")
            return False
    
    def _repair_pdf_for_conversion(self, pdf_path: str) -> str:
        """尝试修复PDF文件"""
        try:
            # 创建临时修复文件
            temp_path = tempfile.mktemp(suffix='.pdf')
            self._temp_files.append(temp_path)
            
            # 使用PyMuPDF重新保存PDF
            doc = fitz.open(pdf_path)
            doc.save(temp_path, garbage=4, deflate=True)
            doc.close()
            
            logger.info(f"PDF文件已修复: {temp_path}")
            return temp_path
            
        except Exception as e:
            logger.error(f"PDF修复失败: {e}")
            return pdf_path

# 转换器元数据
PDF_TO_DOCX_METADATA = ConverterMetadata(
    name="pdf_to_docx",
    description="将PDF文件转换为Word文档格式，支持多种转换方法",
    version="1.0.0",
    author="PDF Converter Team",
    supported_input_formats=["pdf"],
    supported_output_formats=["docx"],
    dependencies=["PyMuPDF", "pdf2docx", "python-docx"],
    priority=10
)