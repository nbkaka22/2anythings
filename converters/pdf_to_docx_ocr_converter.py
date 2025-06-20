# -*- coding: utf-8 -*-
"""PDF转DOCX OCR转换器

使用PyMuPDF + PaddleOCR(主要) + EasyOCR(备用) + pdfplumber组合实现扫描版PDF的文本和图片提取
"""

import os
import fitz  # PyMuPDF
import easyocr
import pdfplumber
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import io
import tempfile
import logging
from typing import Dict, Any, List, Tuple
import numpy as np
import gc
import psutil
from typing import Optional
from pdf2image import convert_from_path
from io import BytesIO
import json

# 导入新的改进组件
from converters.ocr_parameter_adapter import OCRParameterAdapter
from converters.ocr_health_checker import OCRHealthChecker
from converters.performance_monitor import monitor_performance
from converters.retry_manager import retry_ocr_init, retry_file_processing, CircuitBreaker
from converters.enhanced_ocr_methods import EnhancedOCRMethods
from converters.cache_manager import cache_manager
from config.config_validator import config_validator

# PaddleOCR导入（可选）
try:
    from paddleocr import PaddleOCR
    PADDLEOCR_AVAILABLE = True
except ImportError:
    PADDLEOCR_AVAILABLE = False
    PaddleOCR = None

# 导入基类
from converters.converter_interface import ConverterInterface, ConverterMetadata

logger = logging.getLogger('pdf_converter')

class PDFToDocxOCRConverter(ConverterInterface):
    """PDF转DOCX OCR转换器 - 增强版
    
    使用PyMuPDF + PaddleOCR(主要) + EasyOCR(备用) + pdfplumber组合处理扫描版PDF
    """
    
    def __init__(self, config_path: Optional[str] = None):
        self.logger = logging.getLogger(__name__)
        self._temp_files = []
        self._paddle_ocr = None
        self._easy_ocr_reader = None
        self._ocr_engine = 'paddle'  # 默认使用PaddleOCR
        self._force_cpu_mode = False  # 强制CPU模式标志
        
        # 初始化改进组件
        self.parameter_adapter = OCRParameterAdapter()
        self.health_checker = OCRHealthChecker()
        self.circuit_breaker = CircuitBreaker(
            failure_threshold=3,
            recovery_timeout=30.0,
            expected_exception=(RuntimeError, ImportError, OSError)
        )
        self.enhanced_ocr = EnhancedOCRMethods(self)
        
        # 加载配置
        self.config = self._load_config(config_path)
        
        # 验证和修复配置
        self._validate_and_fix_config()
        
        # GPU可用性检查
        self.gpu_available = self._check_gpu_availability()
        
        # 启动健康监控
        self.health_checker.start_monitoring()
        
        # OCR版本信息
        self.ocr_versions = self._detect_ocr_versions()
        
        # 初始化缓存
        self._initialize_cache()
    
    def _check_gpu_availability(self) -> bool:
        """检查GPU可用性"""
        try:
            import torch
            return torch.cuda.is_available()
        except ImportError:
            return False
    
    def _detect_ocr_versions(self) -> Dict[str, str]:
        """检测OCR引擎版本"""
        versions = {}
        
        try:
            import paddleocr
            versions['paddleocr'] = getattr(paddleocr, '__version__', '2.6.0')
        except ImportError:
            versions['paddleocr'] = 'not_installed'
        
        try:
            import easyocr
            versions['easyocr'] = getattr(easyocr, '__version__', '1.6.0')
        except ImportError:
            versions['easyocr'] = 'not_installed'
        
        self.logger.info(f"检测到OCR版本: {versions}")
        return versions
    
    def _should_force_cpu_mode(self, exception: Exception) -> bool:
        """判断是否应该强制使用CPU模式"""
        error_msg = str(exception).lower()
        gpu_error_keywords = [
            'cuda', 'gpu', 'memory', 'device', 'cudnn', 
            'out of memory', 'nvidia', 'driver'
        ]
        return any(keyword in error_msg for keyword in gpu_error_keywords)
    
    def _load_config(self, config_path: Optional[str] = None) -> Dict[str, Any]:
        """加载配置文件"""
        default_config = {
            "gpu_memory_optimization": {
                "enable_monitoring": True,
                "memory_threshold": 85.0,
                "cleanup_interval": 30,
                "oom_handling": True
            },
            "ocr_optimization": {
                "batch_size": 4,
                "max_image_size": 2048,
                "confidence_threshold": 0.5,
                "paddleocr_device_mode": "auto",
                "use_legacy_api": False,
                "api_compatibility_mode": True
            },
            "fallback_strategy": {
                "paddle_to_easy": True,
                "gpu_to_cpu": True,
                "auto_switch": True
            },
            "logging": {
                "level": "INFO",
                "performance_logging": True,
                "error_details": True
            }
        }
        
        if config_path and os.path.exists(config_path):
            try:
                with open(config_path, 'r', encoding='utf-8') as f:
                    user_config = json.load(f)
                    # 合并配置
                    default_config.update(user_config)
            except Exception as e:
                    self.logger.warning(f"加载配置文件失败: {e}，使用默认配置")
        
        return default_config
        
    def _validate_and_fix_config(self):
        """验证和修复配置"""
        try:
            # 如果有配置文件路径，验证配置文件
            if hasattr(self, 'config_path') and self.config_path:
                validation_result = config_validator.validate_gpu_memory_config(self.config_path)
                
                if not validation_result['is_valid']:
                    self.logger.warning(f"配置验证失败，发现 {len(validation_result['errors'])} 个错误")
                    
                    # 尝试修复配置
                    if validation_result['config']:
                        fixed_config = config_validator.fix_config_issues(validation_result['config'])
                        self.config.update(fixed_config)
                        
                        # 保存修复后的配置
                        if config_validator.save_config(fixed_config, self.config_path):
                            self.logger.info("配置已自动修复并保存")
                    
                if validation_result['warnings']:
                    self.logger.warning(f"配置验证发现 {len(validation_result['warnings'])} 个警告")
                    for warning in validation_result['warnings']:
                        self.logger.warning(f"配置警告: {warning}")
            
        except Exception as e:
            self.logger.error(f"配置验证失败: {e}")
    
    def _initialize_cache(self):
        """初始化缓存"""
        try:
            # 设置缓存配置
            cache_config = self.config.get('cache', {})
            max_cache_size = cache_config.get('max_size_mb', 500)
            
            # 配置缓存管理器
            cache_manager.max_cache_size_mb = max_cache_size
            
            # 优化缓存
            optimization_result = cache_manager.optimize_cache()
            if optimization_result.get('success'):
                self.logger.info(f"缓存初始化完成，当前大小: {optimization_result.get('final_size_mb', 0)}MB")
            
            # 记录缓存统计
            cache_stats = cache_manager.get_cache_stats()
            self.logger.debug(f"缓存统计: {cache_stats}")
            
        except Exception as e:
            self.logger.error(f"缓存初始化失败: {e}")
        
    @retry_ocr_init
    @monitor_performance(track_memory=True, track_gpu=True)
    def _get_paddle_ocr(self):
        """获取PaddleOCR实例（延迟初始化）- 增强版"""
        if self._paddle_ocr is None and PADDLEOCR_AVAILABLE:
            try:
                # 获取PaddleOCR版本
                paddle_version = self.ocr_versions.get('paddleocr', '2.6.0')
                
                # 确定设备模式
                gpu_available = self._check_gpu_support()
                device_available = gpu_available and not self._force_cpu_mode
                
                # 如果GPU可用，先清理GPU内存
                if device_available:
                    self._clear_gpu_memory()
                
                # 基础参数
                base_params = {
                    'use_angle_cls': True,
                    'lang': 'ch',
                    'show_log': False,
                    'gpu_available': device_available
                }
                
                # 使用参数适配器获取兼容参数
                adapted_params = self.parameter_adapter.adapt_paddleocr_parameters(
                    base_params, paddle_version, device_available
                )
                
                logger.info(f"初始化PaddleOCR v{paddle_version}，参数: {adapted_params}")
                
                # 使用熔断器保护初始化，添加超时机制
                @self.circuit_breaker
                def init_paddle():
                    import signal
                    import threading
                    
                    result = [None]
                    exception = [None]
                    
                    def target():
                        try:
                            result[0] = PaddleOCR(**adapted_params)
                        except Exception as e:
                            exception[0] = e
                    
                    # 使用线程和超时机制
                    thread = threading.Thread(target=target)
                    thread.daemon = True
                    thread.start()
                    thread.join(timeout=60)  # 60秒超时
                    
                    if thread.is_alive():
                        logger.error("PaddleOCR初始化超时（60秒），可能是GPU初始化问题")
                        # 强制使用CPU模式重试
                        adapted_params_cpu = adapted_params.copy()
                        if 'use_gpu' in adapted_params_cpu:
                            adapted_params_cpu['use_gpu'] = False
                        if 'device' in adapted_params_cpu:
                            adapted_params_cpu['device'] = 'cpu'
                        logger.info("尝试使用CPU模式初始化PaddleOCR")
                        return PaddleOCR(**adapted_params_cpu)
                    
                    if exception[0]:
                        raise exception[0]
                    
                    return result[0]

                self._paddle_ocr = init_paddle()
                logger.info("PaddleOCR初始化成功")
                
                # 更新健康检查统计
                self.health_checker.update_stats(0.5, success=True)
                        
            except Exception as e:
                logger.error(f"PaddleOCR初始化失败: {e}")
                
                # 更新健康检查统计
                self.health_checker.update_stats(0.5, success=False)
                
                # 智能错误处理
                if self._should_force_cpu_mode(e):
                    logger.warning("检测到GPU相关错误，强制使用CPU模式")
                    self._force_cpu_mode = True
                    return self._get_paddle_ocr()  # 递归调用，使用CPU模式
                
                # 如果PaddleOCR完全失败，返回None，将使用EasyOCR作为备用
                logger.warning("PaddleOCR不可用，将使用EasyOCR作为备用")
                self._paddle_ocr = None
                self._ocr_engine = 'easy'
                return None
                
        return self._paddle_ocr
    
    def _get_easy_ocr_reader(self):
        """获取EasyOCR识别器实例（延迟初始化）"""
        if self._easy_ocr_reader is None:
            try:
                # 检查GPU可用性，如果之前有内存错误则强制CPU模式
                gpu_available = self._check_gpu_support() and not getattr(self, '_force_cpu_mode', False)
                
                # 如果GPU可用，先清理GPU内存
                if gpu_available:
                    self._clear_gpu_memory()
                
                # 初始化EasyOCR，支持中英文
                self._easy_ocr_reader = easyocr.Reader(['ch_sim', 'en'], gpu=gpu_available)
                
                if gpu_available:
                    logger.info("EasyOCR初始化成功 (GPU模式)")
                else:
                    logger.info("EasyOCR初始化成功 (CPU模式)")
                    
            except Exception as e:
                # 如果GPU模式失败，尝试CPU模式
                if gpu_available:
                    logger.warning(f"EasyOCR GPU模式初始化失败，尝试CPU模式: {e}")
                    try:
                        self._clear_gpu_memory()
                        self._easy_ocr_reader = easyocr.Reader(['ch_sim', 'en'], gpu=False)
                        self._force_cpu_mode = True  # 标记强制CPU模式
                        logger.info("EasyOCR初始化成功 (CPU模式)")
                    except Exception as e2:
                        logger.error(f"EasyOCR初始化失败: {e2}")
                        raise
                else:
                    logger.error(f"EasyOCR初始化失败: {e}")
                    raise
        return self._easy_ocr_reader
    
    def _check_gpu_support(self) -> bool:
        """检查GPU支持情况"""
        try:
            import torch
            if torch.cuda.is_available():
                # 检查GPU内存使用情况
                gpu_memory_info = self._get_gpu_memory_info()
                if gpu_memory_info:
                    free_memory_gb = gpu_memory_info['free'] / (1024**3)
                    if free_memory_gb < 1.0:  # 如果可用内存少于1GB，使用CPU模式
                        logger.warning(f"GPU可用内存不足 ({free_memory_gb:.2f}GB)，将使用CPU模式")
                        return False
                
                logger.info(f"检测到CUDA GPU: {torch.cuda.get_device_name(0)}")
                return True
            else:
                logger.info("未检测到CUDA GPU，将使用CPU模式")
                return False
        except ImportError:
            logger.info("PyTorch未安装，将使用CPU模式")
            return False
        except Exception as e:
            logger.warning(f"GPU检查失败: {e}，将使用CPU模式")
            return False
    
    def _get_gpu_memory_info(self) -> Dict[str, int]:
        """获取GPU内存信息"""
        try:
            import torch
            if torch.cuda.is_available():
                total_memory = torch.cuda.get_device_properties(0).total_memory
                allocated_memory = torch.cuda.memory_allocated(0)
                free_memory = total_memory - allocated_memory
                return {
                    'total': total_memory,
                    'allocated': allocated_memory,
                    'free': free_memory
                }
        except Exception as e:
            logger.warning(f"获取GPU内存信息失败: {e}")
        return None
    
    def _clear_gpu_memory(self):
        """清理GPU内存"""
        try:
            import torch
            if torch.cuda.is_available():
                torch.cuda.empty_cache()
                torch.cuda.synchronize()
                logger.debug("GPU内存已清理")
        except Exception as e:
            logger.warning(f"清理GPU内存失败: {e}")
    
    @property
    def name(self) -> str:
        return "pdf_to_docx_ocr"
    
    @property
    def description(self) -> str:
        return "使用OCR技术将扫描版PDF转换为Word文档，支持PaddleOCR(主要)+EasyOCR(备用)双引擎文本识别和图片提取"
    
    @property
    def supported_input_formats(self) -> List[str]:
        return ["pdf"]
    
    @property
    def supported_output_formats(self) -> List[str]:
        return ["docx"]
    
    @property
    def version(self) -> str:
        return "1.0.0"
    
    def validate_input(self, input_path: str) -> bool:
        """验证PDF文件是否有效"""
        try:
            if not os.path.exists(input_path):
                logger.error(f"文件不存在: {input_path}")
                return False
            
            if not input_path.lower().endswith('.pdf'):
                logger.error(f"文件格式不正确，需要PDF文件: {input_path}")
                return False
            
            # 尝试打开PDF文件
            doc = fitz.open(input_path)
            if doc.page_count == 0:
                logger.error(f"PDF文件为空: {input_path}")
                doc.close()
                return False
            
            doc.close()
            return True
            
        except Exception as e:
            logger.error(f"PDF文件验证失败: {e}")
            return False
    
    def _detect_pdf_type(self, pdf_path: str) -> str:
        """检测PDF类型：文本型或扫描型"""
        try:
            # 使用pdfplumber检测文本内容
            with pdfplumber.open(pdf_path) as pdf:
                text_pages = 0
                total_pages = len(pdf.pages)
                
                # 检查前几页是否有文本
                check_pages = min(3, total_pages)
                for i in range(check_pages):
                    page = pdf.pages[i]
                    text = page.extract_text()
                    if text and text.strip():
                        text_pages += 1
                
                # 如果超过一半的检查页面有文本，认为是文本型PDF
                if text_pages >= check_pages * 0.5:
                    return "text"
                else:
                    return "scanned"
                    
        except Exception as e:
            logger.warning(f"PDF类型检测失败: {e}，默认为扫描型")
            return "scanned"
    
    def _extract_images_from_page(self, pdf_doc, page_num: int) -> List[Tuple[Image.Image, Dict]]:
        """从PDF页面提取图片"""
        images = []
        try:
            page = pdf_doc[page_num]
            image_list = page.get_images()
            
            for img_index, img in enumerate(image_list):
                try:
                    xref = img[0]
                    base_image = pdf_doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    image_ext = base_image["ext"]
                    
                    # 转换为PIL图像
                    pil_image = Image.open(io.BytesIO(image_bytes))
                    
                    # 获取图像在页面中的位置信息
                    image_info = {
                        'index': img_index,
                        'ext': image_ext,
                        'width': base_image["width"],
                        'height': base_image["height"]
                    }
                    
                    images.append((pil_image, image_info))
                    
                except Exception as e:
                    logger.warning(f"提取第{page_num+1}页第{img_index+1}个图像失败: {e}")
                    
        except Exception as e:
            logger.error(f"提取第{page_num+1}页图像失败: {e}")
            
        return images
    
    def _convert_page_to_image(self, pdf_doc, page_num: int, dpi: int = 300) -> Image.Image:
        """将PDF页面转换为高分辨率图像"""
        try:
            page = pdf_doc[page_num]
            # 设置缩放比例以获得指定DPI
            zoom = dpi / 72.0
            mat = fitz.Matrix(zoom, zoom)
            
            # 渲染页面为图像
            pix = page.get_pixmap(matrix=mat)
            img_data = pix.tobytes("png")
            
            # 转换为PIL图像
            pil_image = Image.open(io.BytesIO(img_data))
            return pil_image
            
        except Exception as e:
            logger.error(f"页面转图像失败: {e}")
            raise
    
    def _ocr_image(self, image: Image.Image) -> List[Dict]:
        """对图像进行OCR识别
        
        Args:
            image: PIL图像对象
            
        Returns:
            包含文本信息的字典列表，每个字典包含bbox、text、confidence
        """
        # 优先使用PaddleOCR（除非被强制使用CPU模式且内存不足）
        if self._ocr_engine == 'paddle' and not getattr(self, '_force_cpu_mode', False):
            try:
                return self._ocr_with_paddle(image)
            except Exception as e:
                logger.warning(f"PaddleOCR识别失败: {e}，切换到EasyOCR")
                self._ocr_engine = 'easy'  # 切换到备用引擎
                # 如果是内存错误，清理GPU内存
                if "CUDA" in str(e) or "memory" in str(e).lower():
                    self._clear_gpu_memory()
                    self._force_cpu_mode = True
        
        # 使用EasyOCR作为备用方案
        try:
            return self._ocr_with_easy(image)
        except Exception as e:
            # 如果是内存错误，尝试清理内存后重试一次
            if "CUDA" in str(e) or "memory" in str(e).lower():
                logger.warning(f"EasyOCR内存错误，清理内存后重试: {e}")
                self._clear_gpu_memory()
                self._force_cpu_mode = True
                # 重新初始化EasyOCR为CPU模式
                self._easy_ocr_reader = None
                try:
                    return self._ocr_with_easy(image)
                except Exception as e2:
                    logger.error(f"重试后仍然失败: {e2}")
            else:
                logger.error(f"所有OCR引擎都失败: {e}")
            return []
    
    def _ocr_with_paddle(self, image: Image.Image) -> List[Dict]:
        """使用PaddleOCR进行文字识别"""
        try:
            # 检查GPU内存状态
            gpu_memory_info = self._get_gpu_memory_info()
            if gpu_memory_info:
                free_memory_gb = gpu_memory_info['free'] / (1024**3)
                if free_memory_gb < 0.5:  # 如果可用内存少于0.5GB，先清理
                    logger.warning(f"GPU内存不足 ({free_memory_gb:.2f}GB)，清理内存")
                    self._clear_gpu_memory()
            
            # 转换PIL图像为numpy数组
            try:
                img_array = np.array(image, dtype=np.uint8)
            except Exception as np_error:
                logger.warning(f"NumPy数组转换失败: {np_error}，尝试其他方法")
                # 备用方法：保存为临时文件
                import tempfile
                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_file:
                    image.save(tmp_file.name)
                    img_array = tmp_file.name
                    self._temp_files.append(tmp_file.name)  # 记录临时文件用于清理
            
            # 使用PaddleOCR进行文字识别
            paddle_ocr = self._get_paddle_ocr()
            if paddle_ocr is None:
                raise Exception("PaddleOCR未初始化")
            
            # PaddleOCR 3.0+ 使用predict方法，旧版本使用ocr方法，添加超时保护
            import threading
            
            def call_paddle_ocr_with_timeout():
                result_container = [None]
                exception_container = [None]
                
                def target():
                    try:
                        # 尝试使用新的predict方法
                        result_container[0] = paddle_ocr.predict(img_array)
                    except (AttributeError, TypeError) as e:
                        # 如果predict方法不存在或参数不兼容，尝试旧的ocr方法
                        logger.warning(f"predict方法调用失败，尝试ocr方法: {e}")
                        try:
                            result_container[0] = paddle_ocr.ocr(img_array, cls=True)
                        except Exception as e2:
                            # 如果cls参数不支持，尝试不带cls参数
                            logger.warning(f"ocr方法(cls=True)调用失败，尝试无参数: {e2}")
                            result_container[0] = paddle_ocr.ocr(img_array)
                    except Exception as e:
                        exception_container[0] = e
                
                thread = threading.Thread(target=target)
                thread.daemon = True
                thread.start()
                thread.join(timeout=30)  # 30秒超时
                
                if thread.is_alive():
                    logger.error("PaddleOCR处理超时（30秒），跳过此图像")
                    return None
                
                if exception_container[0]:
                    raise exception_container[0]
                
                return result_container[0]
            
            results = call_paddle_ocr_with_timeout()
            if results is None:
                logger.warning("PaddleOCR处理超时，返回空结果")
                return []
            
            # 处理PaddleOCR结果
            ocr_data = []
            if results:
                # 处理不同版本PaddleOCR的返回格式
                if isinstance(results, list) and len(results) > 0 and results[0]:
                    # 传统ocr方法返回格式: [[[bbox], (text, confidence)]]
                    for line in results[0]:
                        if line and len(line) == 2:
                            bbox, (text, confidence) = line
                            if confidence > 0.5:  # 只保留置信度较高的结果
                                # 转换bbox格式以兼容EasyOCR格式
                                formatted_bbox = [[bbox[0][0], bbox[0][1]], [bbox[1][0], bbox[1][1]], 
                                                [bbox[2][0], bbox[2][1]], [bbox[3][0], bbox[3][1]]]
                                ocr_data.append({
                                    'bbox': formatted_bbox,
                                    'text': text,
                                    'confidence': confidence
                                })
                elif isinstance(results, dict):
                    # predict方法返回字典格式，包含res字段
                    logger.info(f"检测到predict方法返回格式: {type(results)}")
                    if 'res' in results and 'dt_polys' in results['res'] and 'rec_texts' in results['res']:
                        dt_polys = results['res']['dt_polys']
                        rec_texts = results['res']['rec_texts']
                        # 如果有置信度信息
                        rec_scores = results['res'].get('rec_scores', [1.0] * len(rec_texts))
                        
                        for i, (poly, text) in enumerate(zip(dt_polys, rec_texts)):
                            confidence = rec_scores[i] if i < len(rec_scores) else 1.0
                            if confidence > 0.5:  # 只保留置信度较高的结果
                                # 转换bbox格式以兼容EasyOCR格式
                                formatted_bbox = [[poly[0][0], poly[0][1]], [poly[1][0], poly[1][1]], 
                                                [poly[2][0], poly[2][1]], [poly[3][0], poly[3][1]]]
                                ocr_data.append({
                                    'bbox': formatted_bbox,
                                    'text': text,
                                    'confidence': confidence
                                })
                else:
                    logger.warning(f"未知的返回格式: {type(results)}, 内容: {results}")
            
            logger.debug(f"PaddleOCR识别到 {len(ocr_data)} 个文本块")
            return ocr_data
            
        except Exception as e:
            # 如果是内存错误，清理GPU内存
            if "CUDA" in str(e) or "memory" in str(e).lower():
                self._clear_gpu_memory()
            logger.error(f"PaddleOCR识别失败: {e}")
            raise
    
    def _ocr_with_easy(self, image: Image.Image) -> List[Dict]:
        """使用EasyOCR进行文字识别"""
        try:
            # 检查GPU内存状态（如果未强制CPU模式）
            if not getattr(self, '_force_cpu_mode', False):
                gpu_memory_info = self._get_gpu_memory_info()
                if gpu_memory_info:
                    free_memory_gb = gpu_memory_info['free'] / (1024**3)
                    if free_memory_gb < 0.5:  # 如果可用内存少于0.5GB，先清理
                        logger.warning(f"GPU内存不足 ({free_memory_gb:.2f}GB)，清理内存")
                        self._clear_gpu_memory()
            
            # 转换PIL图像为numpy数组
            try:
                img_array = np.array(image, dtype=np.uint8)
            except Exception as np_error:
                logger.warning(f"NumPy数组转换失败: {np_error}，尝试其他方法")
                # 备用方法：转换为RGB模式
                if image.mode != 'RGB':
                    image = image.convert('RGB')
                img_array = np.asarray(image, dtype=np.uint8)
            
            # 使用EasyOCR进行文字识别
            easy_ocr_reader = self._get_easy_ocr_reader()
            results = easy_ocr_reader.readtext(img_array)
            
            # 处理EasyOCR结果
            ocr_data = []
            for (bbox, text, confidence) in results:
                if confidence > 0.5:  # 只保留置信度较高的结果
                    ocr_data.append({
                        'bbox': bbox,
                        'text': text,
                        'confidence': confidence
                    })
            
            logger.debug(f"EasyOCR识别到 {len(ocr_data)} 个文本块")
            return ocr_data
            
        except Exception as e:
            # 如果是内存错误，清理GPU内存
            if "CUDA" in str(e) or "memory" in str(e).lower():
                self._clear_gpu_memory()
            logger.error(f"EasyOCR识别失败: {e}")
            raise
    
    def _analyze_layout(self, ocr_data: List[Dict], image_width: int, image_height: int) -> List[Dict]:
        """分析文本布局，重构段落结构"""
        if not ocr_data:
            return []
        
        # 按Y坐标排序（从上到下）
        sorted_data = sorted(ocr_data, key=lambda x: min([point[1] for point in x['bbox']]))
        
        # 分组为行
        lines = []
        current_line = []
        current_y = None
        y_threshold = image_height * 0.02  # Y坐标阈值
        
        for item in sorted_data:
            bbox = item['bbox']
            y_center = sum([point[1] for point in bbox]) / 4
            
            if current_y is None or abs(y_center - current_y) <= y_threshold:
                current_line.append(item)
                current_y = y_center if current_y is None else (current_y + y_center) / 2
            else:
                if current_line:
                    lines.append(current_line)
                current_line = [item]
                current_y = y_center
        
        if current_line:
            lines.append(current_line)
        
        # 每行内按X坐标排序（从左到右）
        for line in lines:
            line.sort(key=lambda x: min([point[0] for point in x['bbox']]))
        
        # 合并为段落
        paragraphs = []
        for line in lines:
            line_text = ' '.join([item['text'] for item in line])
            if line_text.strip():
                paragraphs.append({
                    'text': line_text.strip(),
                    'y_position': min([point[1] for point in line[0]['bbox']])
                })
        
        return paragraphs
    
    @monitor_performance(track_memory=True, track_gpu=True)
    @retry_file_processing
    def convert(self, input_path: str, output_path: str, **kwargs) -> bool:
        """执行PDF到DOCX的OCR转换 - 增强版"""
        try:
            logger.info(f"开始OCR转换: {input_path} -> {output_path}")
            
            # 健康检查
            health_report = self.health_checker.check_ocr_health()
            if health_report["status"] == "critical":
                self.logger.warning("系统健康状态异常，但继续执行转换")
                for alert in health_report["alerts"]:
                    if alert["level"] == "critical":
                        self.logger.warning(f"健康警报: {alert['message']}")
            
            # 确保输出目录存在
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            # 检测PDF类型
            pdf_type = self._detect_pdf_type(input_path)
            logger.info(f"PDF类型: {pdf_type}")
            
            # 打开PDF文档
            pdf_doc = fitz.open(input_path)
            total_pages = pdf_doc.page_count
            
            # 创建Word文档
            word_doc = Document()
            
            # 批处理配置
            batch_size = self.config["ocr_optimization"]["batch_size"]
            
            # 收集需要OCR处理的图像
            images_to_process = []
            page_mappings = []
            
            # 处理每一页
            for page_num in range(total_pages):
                logger.info(f"处理第{page_num + 1}/{total_pages}页")
                
                try:
                    if page_num > 0:
                        word_doc.add_page_break()
                    
                    # 添加页面标题
                    title_para = word_doc.add_paragraph(f"第{page_num + 1}页")
                    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    title_run = title_para.runs[0]
                    title_run.font.size = Pt(14)
                    title_run.bold = True
                    
                    if pdf_type == "scanned":
                        # 扫描版PDF：收集图像用于批量OCR处理
                        try:
                            # 将页面转换为图像
                            page_image = self._convert_page_to_image(pdf_doc, page_num)
                            images_to_process.append(page_image)
                            page_mappings.append({
                                'page_num': page_num,
                                'word_doc_position': len(word_doc.paragraphs)
                            })
                            
                            # 提取并添加图片
                            images = self._extract_images_from_page(pdf_doc, page_num)
                            for img, img_info in images:
                                try:
                                    # 保存图片到临时文件
                                    temp_img_path = tempfile.mktemp(suffix=f".{img_info['ext']}")
                                    self._temp_files.append(temp_img_path)
                                    img.save(temp_img_path)
                                    
                                    # 添加到Word文档
                                    word_doc.add_paragraph()  # 空行
                                    para = word_doc.add_paragraph()
                                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    run = para.runs[0] if para.runs else para.add_run()
                                    
                                    # 计算合适的图片尺寸
                                    max_width = Inches(6)
                                    if img.width > img.height:
                                        width = max_width
                                        height = max_width * img.height / img.width
                                    else:
                                        height = max_width
                                        width = max_width * img.width / img.height
                                    
                                    run.add_picture(temp_img_path, width=width, height=height)
                                    
                                except Exception as e:
                                    logger.warning(f"添加图片失败: {e}")
                                    word_doc.add_paragraph(f"[图片添加失败: {e}]")
                        
                        except Exception as e:
                            logger.error(f"处理第{page_num + 1}页失败: {e}")
                            word_doc.add_paragraph(f"[页面处理失败: {e}]")
                    
                    else:
                        # 文本型PDF：直接提取文本
                        try:
                            page = pdf_doc[page_num]
                            text = page.get_text()
                            
                            if text.strip():
                                # 按段落分割文本
                                paragraphs = text.split('\n\n')
                                for paragraph in paragraphs:
                                    if paragraph.strip():
                                        word_doc.add_paragraph(paragraph.strip())
                            else:
                                word_doc.add_paragraph("[未检测到文本内容]")
                                
                            # 提取图片
                            images = self._extract_images_from_page(pdf_doc, page_num)
                            for img, img_info in images:
                                try:
                                    temp_img_path = tempfile.mktemp(suffix=f".{img_info['ext']}")
                                    self._temp_files.append(temp_img_path)
                                    img.save(temp_img_path)
                                    
                                    word_doc.add_paragraph()
                                    para = word_doc.add_paragraph()
                                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    run = para.runs[0] if para.runs else para.add_run()
                                    
                                    max_width = Inches(6)
                                    if img.width > img.height:
                                        width = max_width
                                        height = max_width * img.height / img.width
                                    else:
                                        height = max_width
                                        width = max_width * img.width / img.height
                                    
                                    run.add_picture(temp_img_path, width=width, height=height)
                                    
                                except Exception as e:
                                    logger.warning(f"添加图片失败: {e}")
                        
                        except Exception as e:
                            logger.error(f"文本提取第{page_num + 1}页失败: {e}")
                            word_doc.add_paragraph(f"[文本提取失败: {e}]")
                    
                    # 定期内存清理
                    if (page_num + 1) % batch_size == 0:
                        self._clear_gpu_memory()
                        gc.collect()
                        
                except Exception as page_error:
                    logger.error(f"处理第{page_num + 1}页失败: {page_error}")
                    word_doc.add_paragraph(f"[第{page_num + 1}页 - 处理失败: {str(page_error)}]")
                    continue
            
            # 批量OCR处理收集的图像
            if images_to_process and pdf_type == "scanned":
                logger.info(f"开始批量OCR处理 {len(images_to_process)} 张图像")
                try:
                    # 使用增强OCR方法进行批量处理
                    batch_results = self.enhanced_ocr.batch_process_images(
                        images_to_process
                    )
                    
                    # 将OCR结果插入到对应的Word文档位置
                    for i, text_result in enumerate(batch_results):
                        if i < len(page_mappings):
                            mapping = page_mappings[i]
                            page_num = mapping['page_num']
                            
                            # 记录OCR处理结果
                            logger.info(f"第{page_num + 1}页OCR处理完成，文本长度: {len(text_result)}")
                            
                            if text_result and not text_result.startswith('['):
                                # 直接插入识别的文本
                                para = word_doc.add_paragraph(text_result)
                                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            else:
                                word_doc.add_paragraph("[未识别到文本内容]")
                                
                except Exception as e:
                    logger.error(f"批量OCR处理失败: {e}")
                    # 回退到单独处理
                    for i, page_image in enumerate(images_to_process):
                        try:
                            ocr_data = self._ocr_image(page_image)
                            if ocr_data:
                                paragraphs = self._analyze_layout(ocr_data, page_image.width, page_image.height)
                                for para_data in paragraphs:
                                    para = word_doc.add_paragraph(para_data['text'])
                                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            else:
                                word_doc.add_paragraph("[未识别到文本内容]")
                        except Exception as page_error:
                            logger.error(f"处理图像{i}失败: {page_error}")
                            word_doc.add_paragraph(f"[图像处理失败: {page_error}]")
            
            pdf_doc.close()
            
            # 保存Word文档
            word_doc.save(output_path)
            
            # 最终清理
            self._clear_gpu_memory()
            gc.collect()
            
            # 获取缓存统计信息
            cache_stats = cache_manager.get_cache_stats()
            logger.info(f"缓存统计: 命中率={cache_stats['hit_rate']:.2f}%, "
                       f"总大小={cache_stats['total_size_mb']:.1f}MB")
            
            # 验证输出文件
            if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                logger.info(f"OCR转换成功: {output_path}")
                # 更新健康检查统计
                self.health_checker.update_stats(1.0, success=True)
                return True
            else:
                logger.error("转换完成但输出文件无效")
                return False
                
        except Exception as e:
            logger.error(f"OCR转换失败: {e}")
            # 更新健康检查统计
            self.health_checker.update_stats(0, success=False)
            return False
    
    def get_default_options(self) -> Dict[str, Any]:
        """获取默认转换选项"""
        return {
            'dpi': 300,
            'confidence_threshold': 0.5,
            'extract_images': True,
            'preserve_layout': True
        }
    
    def extract_text_from_image(self, image_path: str) -> str:
        """从图像中提取文本 - 增强版"""
        try:
            # 检查缓存
            cached_result = cache_manager.get_ocr_result(image_path)
            if cached_result:
                logger.debug(f"使用缓存的OCR结果: {image_path}")
                return cached_result
            
            # 加载图像
            image = Image.open(image_path)
            
            # 使用增强OCR方法
            result = self.enhanced_ocr.extract_text_from_image(image)
            
            # 缓存结果
            cache_manager.cache_ocr_result(image_path, result)
            
            return result
            
        except Exception as e:
            logger.error(f"图像文本提取失败: {e}")
            return ""
    
    def cleanup(self):
        """清理临时文件和GPU内存"""
        # 清理临时文件
        for temp_file in self._temp_files:
            try:
                if os.path.exists(temp_file):
                    os.remove(temp_file)
            except Exception as e:
                logger.warning(f"清理临时文件失败: {temp_file}, {e}")
        
        self._temp_files.clear()
        
        # 清理GPU内存
        self._clear_gpu_memory()
        
        # 重置OCR引擎状态
        self._paddle_ocr = None
        self._easy_ocr_reader = None
        self._ocr_engine = 'paddle'
        self._force_cpu_mode = False

# 转换器元数据
PDF_TO_DOCX_OCR_METADATA = ConverterMetadata(
    name="pdf_to_docx_ocr",
    description="使用OCR技术将扫描版PDF转换为Word文档，支持PaddleOCR(主要)+EasyOCR(备用)双引擎文本识别和图片提取",
    version="1.0.0",
    author="PDF Converter Team",
    supported_input_formats=["pdf"],
    supported_output_formats=["docx"],
    dependencies=["PyMuPDF", "PaddleOCR", "EasyOCR", "pdfplumber", "python-docx", "Pillow"],
    priority=5
)