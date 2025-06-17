import os
import sys
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import threading
import fitz  # PyMuPDF
from PIL import Image
import docx
from docx import Document
from pdf2docx import parse
import io
import time
import tempfile
import re
from typing import Dict, List, Optional

# 导入工具模块
from utils import get_resource_path
from pdf_operations import PDFOperations
from scripts.dependency_checker import DependencyChecker, quick_dependency_check

# 导入插件系统
from converters.converter_factory import ConverterFactory
from converters.plugin_manager import get_plugin_manager, initialize_plugins

class PDFConverter:
    def __init__(self):
        # 在创建GUI之前进行依赖检查
        self._check_dependencies_on_startup()
        
        # 初始化插件系统
        self._initialize_plugin_system()
        
        self.root = tk.Tk()
        self.root.title("PDF格式转换工具")
        self.root.geometry("1024x640")
        self.root.resizable(True, True)
        
        # 设置应用图标
        try:
            from PIL import Image, ImageTk
            
            # 尝试多个可能的图标路径
            icon_paths = [
                get_resource_path("assets/icon.ico"),
                os.path.join(os.path.dirname(os.path.abspath(__file__)), "assets", "icon.ico"),
                "assets/icon.ico"
            ]
            
            icon_set = False
            for icon_path in icon_paths:
                if os.path.exists(icon_path):
                    print(f"找到图标文件: {icon_path}")
                    try:
                        # 方法1: 使用iconbitmap
                        self.root.iconbitmap(icon_path)
                        print(f"使用iconbitmap设置图标成功: {icon_path}")
                        
                        # 方法2: 使用PIL加载图标并设置为PhotoImage
                        try:
                            icon_image = Image.open(icon_path)
                            # 调整图标大小为32x32
                            icon_image = icon_image.resize((32, 32), Image.Resampling.LANCZOS)
                            icon_photo = ImageTk.PhotoImage(icon_image)
                            self.root.iconphoto(True, icon_photo)
                            # 保存引用防止被垃圾回收
                            self.icon_photo = icon_photo
                            print("使用PIL方法设置图标成功")
                        except Exception as pil_error:
                            print(f"PIL方法设置图标失败: {pil_error}")
                        
                        icon_set = True
                        break
                    except Exception as icon_error:
                        print(f"设置图标失败 {icon_path}: {icon_error}")
                        continue
            
            if not icon_set:
                print("未找到可用的图标文件")
                
        except Exception as e:
            print(f"设置图标时出错: {e}")
        
        # 初始化PDF操作模块
        self.pdf_operations = PDFOperations(self.root)
        
        # 当前模式：'convert' 或 'operation'
        self.current_mode = 'convert'
        
        self.setup_ui()
        
    def setup_ui(self):
        # 创建主框架
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # 顶部功能按钮区域
        top_buttons_frame = ttk.Frame(main_frame)
        top_buttons_frame.pack(fill=tk.X, pady=(0, 10))
        
        # PDF转换按钮（主要功能）
        self.pdf_convert_btn = ttk.Button(top_buttons_frame, text="PDF转换", 
                                         style="Primary.TButton", width=15,
                                         command=self.switch_to_convert_mode)
        self.pdf_convert_btn.pack(side=tk.LEFT, padx=(0, 10))
        
        # PDF操作按钮（新增功能）
        self.pdf_operation_btn = ttk.Button(top_buttons_frame, text="PDF操作", 
                                           style="Secondary.TButton", width=15,
                                           command=self.switch_to_operation_mode)
        self.pdf_operation_btn.pack(side=tk.LEFT, padx=(0, 10))
        
        # 主内容区域 - 使用三栏布局：左侧功能列表、中间操作区域、右侧日志
        content_frame = ttk.Frame(main_frame)
        content_frame.pack(fill=tk.BOTH, expand=True)
        
        # 左侧功能列表区域
        self.left_frame = ttk.Frame(content_frame, width=180)
        self.left_frame.pack(side=tk.LEFT, fill=tk.Y, padx=(0, 8))
        self.left_frame.pack_propagate(False)
        
        # 功能列表标题
        self.left_title_label = ttk.Label(self.left_frame, text="PDF转其他", font=("Segoe UI", 12, "bold"))
        self.left_title_label.pack(anchor=tk.W, pady=(0, 10))
        
        # 功能按钮列表
        functions = [
            ("📄 文件转Word", "docx"),
            ("🖼 文件转图片", "image"),
            ("📽 文件转PPT", "pptx"),
            ("📊 文件转Excel", "excel"),
            ("📐 PDF转CAD", "cad"),
            ("📝 PDF转TXT", "txt"),
            ("🌐 PDF转HTML", "html"),
            ("📖 PDF转长图", "long_image"),
            ("📚 PDF转电子书", "ebook")
        ]
        
        self.selected_function = tk.StringVar(value="docx")
        self.selected_function.trace('w', self.on_function_change)
        
        # 转换功能按钮容器
        self.convert_functions_frame = ttk.Frame(self.left_frame)
        self.convert_functions_frame.pack(fill=tk.BOTH, expand=True)
        
        for text, value in functions:
            btn = ttk.Radiobutton(self.convert_functions_frame, text=text, variable=self.selected_function, 
                                value=value, style="Function.TRadiobutton")
            btn.pack(anchor=tk.W, pady=2, padx=5)
        
        # PDF操作功能按钮容器
        self.operation_functions_frame = ttk.Frame(self.left_frame)
        
        # PDF操作功能列表
        operation_functions = [
            ("🗑 删除页面", "delete_pages"),
            ("📋 合并PDF", "merge_pdf"),
            ("✂ 分割PDF", "split_pdf"),
            ("🔄 旋转页面", "rotate_pages"),
            ("📏 调整页面", "resize_pages")
        ]
        
        self.selected_operation = tk.StringVar(value="delete_pages")
        self.selected_operation.trace('w', self.on_operation_change)
        
        for text, value in operation_functions:
            btn = ttk.Radiobutton(self.operation_functions_frame, text=text, variable=self.selected_operation, 
                                value=value, style="Function.TRadiobutton")
            btn.pack(anchor=tk.W, pady=2, padx=5)
        
        # 右侧日志区域
        self.log_frame_container = ttk.Frame(content_frame, width=300)
        self.log_frame_container.pack(side=tk.RIGHT, fill=tk.BOTH, padx=(8, 0))
        self.log_frame_container.pack_propagate(False)
        
        # 中间操作区域（压缩空间）
        self.right_frame = ttk.Frame(content_frame)
        self.right_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 8))
        
        # 转换模式的UI容器
        self.convert_mode_frame = ttk.Frame(self.right_frame)
        self.convert_mode_frame.pack(fill=tk.BOTH, expand=True)
        
        # 操作模式的UI容器
        self.operation_mode_frame = ttk.Frame(self.right_frame)
        
        # 设置转换模式UI
        self.setup_convert_mode_ui()
        
        # 文件选择区域
        file_frame = ttk.LabelFrame(self.convert_mode_frame, text="文件选择", padding="15")
        file_frame.pack(fill=tk.X, pady=(0, 10))
        
        # 文件路径输入框
        path_frame = ttk.Frame(file_frame)
        path_frame.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Label(path_frame, text="文件路径:").pack(anchor=tk.W, pady=(0, 5))
        self.file_path_var = tk.StringVar()
        path_entry = ttk.Entry(path_frame, textvariable=self.file_path_var, state="readonly")
        path_entry.pack(fill=tk.X, pady=(0, 10))
        

        
        # 文件操作按钮
        file_buttons_frame = ttk.Frame(file_frame)
        file_buttons_frame.pack(fill=tk.X)
        
        ttk.Button(file_buttons_frame, text="添加文件", command=self.browse_file).pack(side=tk.LEFT, padx=(0, 5))
        ttk.Button(file_buttons_frame, text="添加文件夹", command=self.browse_folder).pack(side=tk.LEFT, padx=(0, 5))
        ttk.Button(file_buttons_frame, text="清空列表", command=self.clear_files).pack(side=tk.LEFT)
        
        # 转换选项区域
        self.options_frame = ttk.LabelFrame(self.convert_mode_frame, text="转换选项", padding="15")
        self.options_frame.pack(fill=tk.X, pady=(0, 10))
        
        # 图片格式选项（针对转图片功能）
        self.image_format_frame = ttk.Frame(self.options_frame)
        self.image_format_frame.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Label(self.image_format_frame, text="图片格式:").pack(anchor=tk.W)
        
        self.image_format_var = tk.StringVar(value="jpg")
        ttk.Radiobutton(self.image_format_frame, text="JPG格式", variable=self.image_format_var, value="jpg").pack(anchor=tk.W, padx=20)
        ttk.Radiobutton(self.image_format_frame, text="PNG格式", variable=self.image_format_var, value="png").pack(anchor=tk.W, padx=20)
        
        # PPT转换选项（针对转PPT功能）
        self.ppt_frame = ttk.Frame(self.options_frame)
        self.ppt_frame.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Label(self.ppt_frame, text="转换PPT方式:").pack(anchor=tk.W)
        
        self.ppt_method_var = tk.StringVar(value="direct")
        ttk.Radiobutton(self.ppt_frame, text="直接转换", variable=self.ppt_method_var, value="direct").pack(anchor=tk.W, padx=20)
        ttk.Radiobutton(self.ppt_frame, text="通过Word转换", variable=self.ppt_method_var, value="via_word").pack(anchor=tk.W, padx=20)
        
        # OCR选项已移除
        
        # 输出目录选择
        output_frame = ttk.LabelFrame(self.convert_mode_frame, text="输出配置", padding="15")
        output_frame.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Label(output_frame, text="输出目录:").pack(anchor=tk.W, pady=(0, 5))
        
        dir_frame = ttk.Frame(output_frame)
        dir_frame.pack(fill=tk.X, pady=(0, 10))
        
        self.output_dir_var = tk.StringVar()
        ttk.Entry(dir_frame, textvariable=self.output_dir_var, width=40).pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 5))
        ttk.Button(dir_frame, text="浏览", command=self.browse_output_dir).pack(side=tk.RIGHT)
        
        # 兼容性变量（保持原有功能）
        self.mode_var = tk.StringVar(value="document")
        self.format_var = tk.StringVar(value="docx")
        
        # 开始转换按钮
        convert_frame = ttk.Frame(output_frame)
        convert_frame.pack(fill=tk.X, pady=(10, 0))
        
        ttk.Button(convert_frame, text="开始转换", command=self.start_conversion, 
                  style="Convert.TButton", width=20).pack(anchor=tk.CENTER)
        
        # 用于跟踪是否上一条日志是进度信息
        self.last_log_was_progress = False
        
        # 初始化时调用一次功能变化处理
        self.on_function_change()
        
        # 设置右侧日志区域
        self.setup_log_area()
        
        # 兼容性变量
        self.ppt_mode_var = tk.StringVar(value="image")
        self.dpi_var = tk.StringVar(value="200")
        
        # 设置样式
        self.setup_styles()
        
        # 初始化为转换模式
        self.switch_to_convert_mode()
    
    def setup_log_area(self):
        """设置右侧日志区域"""
        # 日志区域标题和控制按钮
        log_header_frame = ttk.Frame(self.log_frame_container)
        log_header_frame.pack(fill=tk.X, pady=(0, 5))
        
        ttk.Label(log_header_frame, text="实时日志", font=("Segoe UI", 11, "bold")).pack(side=tk.LEFT)
        
        # 日志控制按钮
        log_controls_frame = ttk.Frame(log_header_frame)
        log_controls_frame.pack(side=tk.RIGHT)
        
        ttk.Button(log_controls_frame, text="清空", command=self.clear_log, width=6).pack(side=tk.RIGHT, padx=(5, 0))
        ttk.Button(log_controls_frame, text="保存", command=self.save_log, width=6).pack(side=tk.RIGHT)
        
        # 日志显示区域 - 设置固定高度，不再expand
        log_display_frame = ttk.LabelFrame(self.log_frame_container, text="", padding="5")
        log_display_frame.pack(fill=tk.X, pady=(0, 0))
        
        # 创建日志文本框和滚动条 - 设置固定高度
        self.log_text = tk.Text(log_display_frame, wrap=tk.WORD, font=("Consolas", 9), 
                               bg="#f8f9fa", fg="#333333", relief="flat", height=24)
        self.log_text.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        scrollbar = ttk.Scrollbar(log_display_frame, command=self.log_text.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.log_text.config(yscrollcommand=scrollbar.set)
        
        # 配置日志文本框的标签样式
        self.log_text.tag_configure("info", foreground="#0066cc")
        self.log_text.tag_configure("success", foreground="#28a745")
        self.log_text.tag_configure("warning", foreground="#ffc107")
        self.log_text.tag_configure("error", foreground="#dc3545")
        self.log_text.tag_configure("timestamp", foreground="#6c757d", font=("Consolas", 8))
        
        # 转换进度区域（放在日志区域下面）
        progress_frame = ttk.LabelFrame(self.log_frame_container, text="转换进度", padding="8")
        progress_frame.pack(fill=tk.X, pady=(10, 15))
        
        self.progress_var = tk.DoubleVar()
        self.progress_bar = ttk.Progressbar(progress_frame, variable=self.progress_var, maximum=100)
        self.progress_bar.pack(fill=tk.X, padx=5, pady=3)
        
        self.status_var = tk.StringVar(value="就绪")
        ttk.Label(progress_frame, textvariable=self.status_var).pack(anchor=tk.W, pady=3)
        
    def clear_log(self):
        """清空日志"""
        self.log_text.delete(1.0, tk.END)
        
    def save_log(self):
        """保存日志到文件"""
        try:
            from tkinter import filedialog
            import datetime
            
            # 默认文件名包含时间戳
            default_name = f"pdf_converter_log_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
            
            file_path = filedialog.asksaveasfilename(
                defaultextension=".txt",
                filetypes=[("文本文件", "*.txt"), ("所有文件", "*.*")],
                initialname=default_name,
                title="保存日志文件"
            )
            
            if file_path:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(self.log_text.get(1.0, tk.END))
                messagebox.showinfo("成功", f"日志已保存到：{file_path}")
        except Exception as e:
            messagebox.showerror("错误", f"保存日志失败：{str(e)}")
        
    def setup_styles(self):
        style = ttk.Style()
        
        # 尝试设置主题
        try:
            style.theme_use("clam")
        except:
            pass
        
        # 主要功能按钮样式
        style.configure("Primary.TButton", 
                       font=("Segoe UI", 11, "bold"),
                       foreground="white",
                       background="#0078d4")
        
        # 次要功能按钮样式
        style.configure("Secondary.TButton", 
                       font=("Segoe UI", 11),
                       foreground="#0078d4",
                       background="white")
        
        # 功能列表单选按钮样式
        style.configure("Function.TRadiobutton", 
                       font=("Segoe UI", 10),
                       padding=(5, 5))
        
        # 转换按钮样式
        style.configure("Convert.TButton", 
                       font=("Segoe UI", 12, "bold"),
                       foreground="white",
                       background="#0078d4",
                       padding=(10, 8))
        
        # 效果查看按钮样式
        style.configure("Effect.TButton", 
                       font=("Segoe UI", 9),
                       foreground="#0078d4",
                       background="#e1f5fe")
        
        # 拖拽区域样式
        style.configure("Drop.TFrame", 
                       relief="solid",
                       borderwidth=2,
                       background="#f8f9fa")
        
        style.configure("Drop.TLabel", 
                       font=("Segoe UI", 11),
                       foreground="#6c757d",
                       background="#f8f9fa")
        
    def update_format_options(self, *args):
        mode = self.mode_var.get()
        output_format = self.format_var.get()
        
        if mode == "document":
            self.image_formats_frame.grid_remove()
            self.document_formats_frame.grid()
            
            # 显示或隐藏PPT转换说明
            if output_format == "pptx":
                self.ppt_options_frame.grid()
            else:
                self.ppt_options_frame.grid_remove()
                
            if not output_format or output_format in ["png", "jpg"]:
                self.format_var.set("docx")
        else:  # image mode
            self.document_formats_frame.grid_remove()
            self.image_formats_frame.grid()
            self.ppt_options_frame.grid_remove()
            
            if not output_format or output_format in ["docx", "pptx", "txt"]:
                self.format_var.set("png")
    
    def browse_file(self):
        filetypes = [("PDF文件", "*.pdf"), ("所有文件", "*.*")]
        file_path = filedialog.askopenfilename(filetypes=filetypes, title="选择PDF文件")
        
        if file_path:
            self.file_path_var.set(file_path)
            # 默认设置输出目录为文件所在目录
            if not self.output_dir_var.get():
                self.output_dir_var.set(os.path.dirname(file_path))
    
    def browse_folder(self):
        folder_path = filedialog.askdirectory(title="选择包含PDF文件的文件夹")
        
        if folder_path:
            self.file_path_var.set(folder_path)
            # 默认设置输出目录为选择的文件夹
            if not self.output_dir_var.get():
                self.output_dir_var.set(folder_path)
    
    def sanitize_filename(self, filename):
        """清理文件名中的非法字符，确保可以创建文件夹"""
        # Windows系统不允许的字符
        invalid_chars = r'<>:"/\|?*'
        # 替换非法字符为下划线
        for char in invalid_chars:
            filename = filename.replace(char, '_')
        
        # 移除开头和结尾的空格和点号
        filename = filename.strip(' .')
        
        # 如果文件名为空或只包含非法字符，使用默认名称
        if not filename:
            filename = 'converted_pdf'
        
        # 限制文件名长度（Windows路径限制）
        if len(filename) > 100:
            filename = filename[:100]
        
        return filename
    
    def browse_output_dir(self):
        output_dir = filedialog.askdirectory(title="选择输出目录")
        
        if output_dir:
            self.output_dir_var.set(output_dir)
    
    def on_function_change(self, *args):
        """当选择的功能发生变化时调用"""
        selected_func = self.selected_function.get()
        
        # 根据选择的功能显示或隐藏相应的转换选项
        if selected_func == "docx":
            # 文件转Word：显示转换选项区域（OCR选项已移除）
            self.options_frame.pack(fill=tk.X, pady=(0, 10))
            self.image_format_frame.pack_forget()
            self.ppt_frame.pack_forget()
        elif selected_func == "image":
            # 文件转图片：显示转换选项区域，只显示图片格式选项
            self.options_frame.pack(fill=tk.X, pady=(0, 10))
            self.image_format_frame.pack(fill=tk.X, pady=(0, 10))
            self.ppt_frame.pack_forget()
        elif selected_func == "pptx":
            # 文件转PPT：显示转换选项区域，只显示PPT转换选项
            self.options_frame.pack(fill=tk.X, pady=(0, 10))
            self.image_format_frame.pack_forget()
            self.ppt_frame.pack(fill=tk.X, pady=(0, 10))
        else:
            # 其他功能：显示转换选项区域，隐藏所有特定选项
            self.options_frame.pack(fill=tk.X, pady=(0, 10))
            self.image_format_frame.pack_forget()
            self.ppt_frame.pack_forget()
    
    def clear_files(self):
        """清空文件列表"""
        self.file_path_var.set("")
        self.log("已清空文件列表")
    
    def log(self, message, update_last_line=False):
        """记录日志信息
        Args:
            message: 日志消息
            update_last_line: 是否更新最后一行（用于进度显示）
        """
        if update_last_line:
            # 只有当上一条日志也是进度信息时才删除上一行
            if self.last_log_was_progress:
                try:
                    # 删除最后一行
                    self.log_text.delete("end-2l", "end-1l")
                except:
                    pass  # 如果删除失败就忽略
            self.log_text.insert(tk.END, f"{message}\n")
            self.last_log_was_progress = True
        else:
            self.log_text.insert(tk.END, f"{message}\n")
            self.last_log_was_progress = False
        self.log_text.see(tk.END)
        self.root.update_idletasks()  # 立即更新UI
    
    def log_progress(self, current, total, message=""):
        """显示进度信息（同行更新）"""
        progress_msg = f"当前处理进度: {current}/{total}"
        if message:
            progress_msg += f" - {message}"
        self.log(progress_msg, update_last_line=True)
    
    def log_step(self, step_name, details=""):
        """显示处理步骤"""
        step_msg = f"📋 {step_name}"
        if details:
            step_msg += f": {details}"
        self.log(step_msg)
    
    def log_error(self, error_msg, full_error=None):
        """显示错误信息"""
        self.log(f"❌ 错误: {error_msg}")
        if full_error:
            self.log(f"详细错误信息: {str(full_error)}")
    
    def log_success(self, message):
        """显示成功信息"""
        self.log(f"✅ {message}")
    
    def log_fallback(self, message):
        """显示备选方案信息"""
        self.log(f"🔄 备选方案: {message}")
    
    def start_conversion(self):
        # 获取输入路径
        input_path = self.file_path_var.get()
        if not input_path:
            messagebox.showerror("错误", "请选择PDF文件或文件夹")
            return
        
        # 获取输出目录
        output_dir = self.output_dir_var.get()
        if not output_dir:
            messagebox.showerror("错误", "请选择输出目录")
            return
        
        # 根据选择的功能确定转换格式和模式
        selected_func = self.selected_function.get()
        
        if selected_func == "docx":
            self.format_var.set("docx")
            self.mode_var.set("document")
        elif selected_func == "image":
            # 使用用户选择的图片格式
            selected_format = self.image_format_var.get()
            self.format_var.set(selected_format)
            self.mode_var.set("image")
        elif selected_func == "pptx":
            if self.ppt_method_var.get() == "via_word":
                self.format_var.set("pptx_via_word")
            else:
                self.format_var.set("pptx")
            self.mode_var.set("document")
        elif selected_func == "txt":
            self.format_var.set("txt")
            self.mode_var.set("document")
        else:
            # 对于其他格式，暂时使用默认的docx
            self.format_var.set("docx")
            self.mode_var.set("document")
            self.log(f"功能 {selected_func} 暂未实现，使用默认的Word转换")
        
        # 确保输出目录存在
        if not os.path.exists(output_dir):
            try:
                os.makedirs(output_dir)
            except Exception as e:
                messagebox.showerror("错误", f"无法创建输出目录: {str(e)}")
                return
        
        # 获取转换选项
        mode = self.mode_var.get()
        output_format = self.format_var.get()
        dpi = int(self.dpi_var.get()) if mode == "image" else None
        
        # 在新线程中执行转换，避免UI冻结
        threading.Thread(target=self.conversion_thread, args=(input_path, output_dir, mode, output_format, dpi), daemon=True).start()
    
    def conversion_thread(self, input_path, output_dir, mode, output_format, dpi):
        try:
            self.status_var.set("正在准备转换...")
            self.progress_var.set(0)
            
            # 确定要处理的文件列表
            pdf_files = []
            if os.path.isfile(input_path) and input_path.lower().endswith('.pdf'):
                pdf_files = [input_path]
            elif os.path.isdir(input_path):
                for file in os.listdir(input_path):
                    if file.lower().endswith('.pdf'):
                        pdf_files.append(os.path.join(input_path, file))
            
            if not pdf_files:
                self.status_var.set("未找到PDF文件")
                messagebox.showinfo("信息", "未找到PDF文件")
                return
            
            total_files = len(pdf_files)
            self.log(f"找到 {total_files} 个PDF文件待转换")
            
            # 检查目标文件是否已存在
            existing_files = []
            for pdf_file in pdf_files:
                base_name = os.path.splitext(os.path.basename(pdf_file))[0]
                
                # 根据转换格式确定输出文件路径
                if output_format == "docx":
                    output_path = os.path.join(output_dir, f"{base_name}.docx")
                elif output_format in ["pptx", "pptx_via_word"]:
                    output_path = os.path.join(output_dir, f"{base_name}.pptx")
                elif output_format == "txt":
                    output_path = os.path.join(output_dir, f"{base_name}.txt")
                elif output_format in ["jpg", "png"]:
                    # 对于图片格式，检查是否存在任何页面文件
                    page_exists = False
                    page_num = 1
                    while True:
                        page_path = os.path.join(output_dir, f"{base_name}_page{page_num}.{output_format}")
                        if os.path.exists(page_path):
                            page_exists = True
                            break
                        page_num += 1
                        if page_num > 1000:  # 防止无限循环
                            break
                    if page_exists:
                        existing_files.append(f"{base_name} (图片文件)")
                    continue
                else:
                    output_path = os.path.join(output_dir, f"{base_name}.{output_format}")
                
                if os.path.exists(output_path):
                    existing_files.append(os.path.basename(output_path))
            
            # 如果有同名文件存在，询问用户是否覆盖
            if existing_files:
                file_list = "\n".join(existing_files)
                message = f"以下文件已存在：\n\n{file_list}\n\n是否要覆盖这些文件？"
                result = messagebox.askyesno("文件已存在", message, icon="warning")
                if not result:
                    self.status_var.set("转换已取消")
                    self.log("用户取消转换操作")
                    return
            
            # 开始转换
            successful = 0
            failed = 0
            
            for i, pdf_file in enumerate(pdf_files):
                file_name = os.path.basename(pdf_file)
                self.status_var.set(f"正在转换 {file_name} ({i+1}/{total_files})")
                # 文件级别进度显示已精简
                
                try:
                    if mode == "document":
                        if output_format == "docx":
                            # OCR选项已移除
                            self.convert_to_docx(pdf_file, output_dir)
                        elif output_format == "pptx":
                            self.convert_to_pptx(pdf_file, output_dir)
                        elif output_format == "pptx_via_word":
                            self.convert_to_pptx_via_word(pdf_file, output_dir)
                        elif output_format == "txt":
                            self.convert_to_txt(pdf_file, output_dir)
                    else:  # image mode
                        self.convert_to_image(pdf_file, output_dir, output_format, dpi)
                    
                    successful += 1
                    # 单文件转换成功日志已精简
                except Exception as e:
                    failed += 1
                    self.log_error(f"转换失败: {file_name}", e)
                
                # 更新进度
                progress = (i + 1) / total_files * 100
                self.progress_var.set(progress)
            
            # 完成
            self.status_var.set(f"转换完成: {successful}成功, {failed}失败")
            self.log_success(f"批量转换完成 - 成功: {successful}, 失败: {failed}")
            messagebox.showinfo("完成", f"转换完成\n成功: {successful}\n失败: {failed}")
            
        except Exception as e:
            self.status_var.set(f"转换过程中发生错误: {str(e)}")
            self.log(f"错误: {str(e)}")
            messagebox.showerror("错误", f"转换过程中发生错误: {str(e)}")
    
    def convert_to_docx(self, pdf_path, output_dir):
        # 获取文件名（不含扩展名）
        base_name = os.path.splitext(os.path.basename(pdf_path))[0]
        # 确保路径使用正确的分隔符
        output_path = os.path.normpath(os.path.join(output_dir, f"{base_name}.docx"))
        
        self.log("开始转换PDF到DOCX")
        
        # 使用工厂模式获取转换器
        if hasattr(self, 'converter_factory') and self.converter_factory:
            try:
                self.log("使用工厂模式获取PDF到DOCX转换器")
                converter = self.converter_factory.get_converter('pdf', 'docx')
                if converter:
                    self.log(f"找到转换器: {converter.name}")
                    success = converter.convert(pdf_path, output_path)
                    if success:
                        self.log_success(f"转换完成: {base_name}.pdf -> {base_name}.docx")
                        self.log(f"输出文件: {output_path}")
                        return output_path
                    else:
                        self.log_error("转换器执行失败", None)
                else:
                    self.log_error("未找到支持PDF到DOCX转换的转换器", None)
            except Exception as e:
                self.log_error(f"工厂模式转换出错: {e}", e)
        
        raise Exception("没有可用的PDF转DOCX转换器")
    
    def _check_pdf_integrity(self, pdf_path):
        """检查PDF文件的完整性，特别是查找可能导致bandwriter错误的问题"""
        try:
            self.log(f"检查PDF文件: {pdf_path}")
            doc = fitz.open(pdf_path)
            
            page_count = len(doc)
            self.log(f"PDF页数: {page_count}")
            
            problematic_pages = []
            zero_dimension_images = []
            
            for page_num in range(page_count):
                page = doc.load_page(page_num)
                
                # 检查页面中的图像
                image_list = page.get_images()
                for img_index, img in enumerate(image_list):
                    try:
                        # 获取图像信息
                        xref = img[0]
                        
                        # 直接从文档获取图像信息，避免创建Pixmap
                        img_dict = doc.extract_image(xref)
                        width = img_dict.get('width', 0)
                        height = img_dict.get('height', 0)
                        
                        # 检查图像尺寸
                        if width == 0 or height == 0:
                            self.log(f"发现零尺寸图像: 页面 {page_num + 1}, 图像 {img_index + 1}, 尺寸: {width}x{height}")
                            zero_dimension_images.append((page_num, img_index, xref))
                            if page_num not in problematic_pages:
                                problematic_pages.append(page_num)
                        
                    except Exception as e:
                        self.log(f"检查图像时出错: 页面 {page_num + 1}, 图像 {img_index + 1} - {str(e)}")
                        if page_num not in problematic_pages:
                            problematic_pages.append(page_num)
            
            doc.close()
            
            if problematic_pages:
                self.log(f"发现问题页面: {problematic_pages}")
                self.log(f"发现 {len(zero_dimension_images)} 个零尺寸图像")
                return False
            else:
                self.log("PDF文件检查通过")
                return True
                
        except Exception as e:
            self.log(f"PDF完整性检查失败: {str(e)}")
            return False
    
    def _repair_pdf_for_conversion(self, pdf_path):
        """修复PDF文件以便转换，专门处理零尺寸图像问题"""
        try:
            self.log("尝试修复PDF文件，移除零尺寸图像...")
            
            # 创建临时文件
            temp_dir = tempfile.mkdtemp()
            temp_pdf_path = os.path.join(temp_dir, "repaired_bandwriter_fix.pdf")
            
            # 打开原PDF
            doc = fitz.open(pdf_path)
            
            # 创建新的PDF文档
            new_doc = fitz.open()
            
            removed_images_count = 0
            
            # 逐页复制，移除零尺寸图像
            for page_num in range(len(doc)):
                try:
                    page = doc.load_page(page_num)
                    rect = page.rect
                    
                    # 创建新页面
                    new_page = new_doc.new_page(width=rect.width, height=rect.height)
                    
                    # 首先复制页面的基本内容（不包括图像）
                    # 获取页面的显示列表，但排除图像
                    page_dict = page.get_text("dict")
                    
                    # 重新插入文本内容
                    for block in page_dict.get("blocks", []):
                        if "lines" in block:  # 文本块
                            for line in block["lines"]:
                                for span in line["spans"]:
                                    text = span.get("text", "")
                                    if text.strip():
                                        bbox = span["bbox"]
                                        font_size = span.get("size", 12)
                                        try:
                                            new_page.insert_text(
                                                (bbox[0], bbox[1]), 
                                                text, 
                                                fontsize=font_size
                                            )
                                        except:
                                            pass  # 忽略插入失败的文本
                    
                    # 处理图像，只保留有效尺寸的图像
                    image_list = page.get_images()
                    for img_index, img in enumerate(image_list):
                        try:
                            xref = img[0]
                            
                            # 检查图像尺寸
                            img_dict = doc.extract_image(xref)
                            width = img_dict.get('width', 0)
                            height = img_dict.get('height', 0)
                            
                            if width > 0 and height > 0:
                                # 图像尺寸正常，保留
                                try:
                                    pix = fitz.Pixmap(doc, xref)
                                    img_rects = page.get_image_rects(xref)
                                    if img_rects:
                                        for img_rect in img_rects:
                                            new_page.insert_image(img_rect, pixmap=pix)
                                    pix = None
                                except:
                                    pass  # 如果插入失败，跳过这个图像
                            else:
                                # 零尺寸图像，跳过
                                removed_images_count += 1
                                self.log(f"移除零尺寸图像: 页面 {page_num + 1}, 图像 {img_index + 1}, 尺寸: {width}x{height}")
                        
                        except Exception as e:
                            self.log(f"处理图像时出错: 页面 {page_num + 1}, 图像 {img_index + 1} - {str(e)}")
                            continue
                    
                except Exception as e:
                    self.log(f"处理页面时出错 {page_num + 1}: {str(e)}")
                    # 即使出错也要创建一个空页面，保持页面数量
                    try:
                        new_doc.new_page()
                    except:
                        pass
            
            # 保存修复后的PDF
            new_doc.save(temp_pdf_path, garbage=4, deflate=True)
            new_doc.close()
            doc.close()
            
            self.log(f"PDF修复完成: 移除了 {removed_images_count} 个零尺寸图像")
            self.log(f"修复后的文件: {temp_pdf_path}")
            return temp_pdf_path
            
        except Exception as e:
            self.log(f"PDF修复失败: {str(e)}")
            return None
    
    def _fix_bandwriter_error(self, pdf_path, output_path):
        """专门修复bandwriter错误"""
        try:
            import fitz
            from docx import Document
            
            self.log("使用PyMuPDF直接提取内容来避免bandwriter错误...")
            
            # 打开PDF
            doc = fitz.open(pdf_path)
            
            # 创建Word文档
            word_doc = Document()
            
            for page_num in range(len(doc)):
                try:
                    page = doc[page_num]
                    
                    # 添加页面标题
                    if page_num > 0:
                        word_doc.add_page_break()
                    
                    word_doc.add_heading(f'页面 {page_num + 1}', level=2)
                    
                    # 提取文本
                    text = page.get_text()
                    if text.strip():
                        # 按段落分割文本
                        paragraphs = text.split('\n\n')
                        for para in paragraphs:
                            if para.strip():
                                word_doc.add_paragraph(para.strip())
                    
                    # 提取表格
                    try:
                        tables = page.find_tables()
                        for table in tables:
                            word_table = word_doc.add_table(rows=len(table.extract()), cols=len(table.extract()[0]))
                            for i, row in enumerate(table.extract()):
                                for j, cell in enumerate(row):
                                    word_table.cell(i, j).text = str(cell) if cell else ''
                    except:
                        self.log(f"页面 {page_num+1} 表格提取失败")
                    
                    self.log(f"页面 {page_num+1} 内容提取完成")
                    
                except Exception as page_error:
                    self.log(f"页面 {page_num+1} 处理失败: {str(page_error)}")
            
            # 保存文档
            word_doc.save(output_path)
            doc.close()
            
            self.log("bandwriter错误修复方法完成")
            return True
            
        except Exception as e:
            self.log(f"bandwriter错误修复失败: {str(e)}")
            return False
    
    def _try_page_by_page_conversion(self, pdf_path, output_path):
        """尝试逐页转换PDF，跳过有问题的页面"""
        try:
            import fitz  # PyMuPDF
            from docx import Document
            
            self.log("开始逐页转换...")
            
            # 打开PDF文件
            pdf_doc = fitz.open(pdf_path)
            total_pages = len(pdf_doc)
            self.log(f"总页数: {total_pages}")
            
            # 创建新的Word文档
            doc = Document()
            
            successful_pages = 0
            
            for page_num in range(total_pages):
                try:
                    self.log_progress(page_num + 1, total_pages, f"处理页面 {page_num + 1}")
                    
                    # 尝试转换单页
                    temp_pdf_path = pdf_path.replace('.pdf', f'_temp_page_{page_num}.pdf')
                    temp_docx_path = output_path.replace('.docx', f'_temp_page_{page_num}.docx')
                    
                    # 创建单页PDF
                    single_page_doc = fitz.open()
                    single_page_doc.insert_pdf(pdf_doc, from_page=page_num, to_page=page_num)
                    single_page_doc.save(temp_pdf_path)
                    single_page_doc.close()
                    
                    # 尝试转换单页
                    parse(temp_pdf_path, temp_docx_path, 
                          multi_processing=False, cpu_count=1)
                    
                    # 如果成功，将内容添加到主文档
                    temp_doc = Document(temp_docx_path)
                    for paragraph in temp_doc.paragraphs:
                        doc.add_paragraph(paragraph.text)
                    
                    successful_pages += 1
                    self.log(f"✓ 页面 {page_num + 1} 转换成功")
                    
                    # 清理临时文件
                    os.remove(temp_pdf_path)
                    os.remove(temp_docx_path)
                    
                except Exception as page_error:
                    self.log(f"✗ 页面 {page_num + 1} 转换失败: {str(page_error)}")
                    # 使用备用方法提取该页文本
                    try:
                        page = pdf_doc[page_num]
                        text = page.get_text()
                        if text.strip():
                            doc.add_paragraph(f"[页面 {page_num + 1} - 文本提取]")
                            doc.add_paragraph(text)
                            successful_pages += 1
                            self.log(f"✓ 页面 {page_num + 1} 文本提取成功")
                    except:
                        self.log(f"✗ 页面 {page_num + 1} 文本提取也失败")
                    
                    # 清理可能存在的临时文件
                    temp_pdf_path = pdf_path.replace('.pdf', f'_temp_page_{page_num}.pdf')
                    temp_docx_path = output_path.replace('.docx', f'_temp_page_{page_num}.docx')
                    if os.path.exists(temp_pdf_path):
                        os.remove(temp_pdf_path)
                    if os.path.exists(temp_docx_path):
                        os.remove(temp_docx_path)
            
            # 保存最终文档
            if successful_pages > 0:
                doc.save(output_path)
                self.log(f"✓ 页面转换完成: {successful_pages}/{total_pages} 页成功")
                pdf_doc.close()
                return True
            else:
                self.log("✗ 所有页面转换都失败")
                pdf_doc.close()
                return False
                
        except Exception as e:
            self.log(f"✗ 页面转换过程出错: {str(e)}")
            return False

    def _convert_to_docx_fallback(self, pdf_path, output_dir):
        """备用的PDF转DOCX方法，使用原来的实现"""
        try:
            # 获取文件名（不含扩展名）
            base_name = os.path.splitext(os.path.basename(pdf_path))[0]
            # 确保路径使用正确的分隔符
            output_path = os.path.normpath(os.path.join(output_dir, f"{base_name}.docx"))
            
            self.log(f"使用备用方法转换: {base_name}.pdf")
        
            # 创建一个新的Word文档
            doc = Document()
            
            # 打开PDF文件
            pdf_document = fitz.open(pdf_path)
            
            # 遍历每一页
            for page_num in range(len(pdf_document)):
                page = pdf_document.load_page(page_num)
                text = page.get_text("text")
                
                # 清理文本中的XML不兼容字符
                cleaned_text = self._clean_text_for_xml(text)
                
                # 添加页码标题
                doc.add_heading(f"Page {page_num + 1}", level=1)
                
                # 添加文本内容（只有在有内容时才添加）
                if cleaned_text.strip():
                    doc.add_paragraph(cleaned_text)
                else:
                    doc.add_paragraph("[此页无可提取的文本内容]")
                
                # 添加分页符（除了最后一页）
                if page_num < len(pdf_document) - 1:
                    doc.add_page_break()
            
            # 保存Word文档
            doc.save(output_path)
            pdf_document.close()
            self.log(f"✓ 备用方法转换成功: {base_name}.pdf -> {base_name}.docx")
            
        except Exception as e:
            self.log(f"✗ 备用方法转换失败: {base_name}.pdf - {str(e)}")
            raise e
    
    def _clean_text_for_xml(self, text):
        """清理文本中的XML不兼容字符"""
        if not text:
            return ""
        
        # 移除NULL字节和控制字符（保留换行符、制表符和回车符）
        # XML 1.0 允许的字符范围：
        # #x9 | #xA | #xD | [#x20-#xD7FF] | [#xE000-#xFFFD] | [#x10000-#x10FFFF]
        valid_chars = []
        for char in text:
            code = ord(char)
            if (code == 0x09 or  # Tab
                code == 0x0A or  # Line feed
                code == 0x0D or  # Carriage return
                (0x20 <= code <= 0xD7FF) or  # Basic multilingual plane
                (0xE000 <= code <= 0xFFFD) or  # Private use area
                (0x10000 <= code <= 0x10FFFF)):  # Supplementary planes
                valid_chars.append(char)
            else:
                # 替换无效字符为空格
                valid_chars.append(' ')
        
        cleaned_text = ''.join(valid_chars)
        
        # 移除多余的空白字符
        cleaned_text = re.sub(r'\s+', ' ', cleaned_text)
        
        return cleaned_text.strip()
    
    def convert_to_txt(self, pdf_path, output_dir):
        # 获取文件名（不含扩展名）
        base_name = os.path.splitext(os.path.basename(pdf_path))[0]
        # 确保路径使用正确的分隔符
        output_path = os.path.normpath(os.path.join(output_dir, f"{base_name}.txt"))
        
        self.log_step("文本转换", f"开始转换 {base_name}.pdf")
        
        # 打开PDF文件
        pdf_document = fitz.open(pdf_path)
        total_pages = len(pdf_document)
        self.log(f"PDF总页数: {total_pages}")
        
        # 创建文本文件
        with open(output_path, 'w', encoding='utf-8') as txt_file:
            # 遍历每一页
            for page_num in range(total_pages):
                self.log_progress(page_num + 1, total_pages, f"提取页面 {page_num + 1} 文本")
                
                page = pdf_document.load_page(page_num)
                text = page.get_text("text")
                
                # 写入页码标题
                txt_file.write(f"===== Page {page_num + 1} =====\n\n")
                
                # 写入文本内容
                txt_file.write(text)
                txt_file.write("\n\n")
        
        pdf_document.close()
        self.log_success(f"文本转换完成: {base_name}.txt")
    
    def convert_to_pptx(self, pdf_path, output_dir):
        """将PDF转换为PowerPoint演示文稿"""
        try:
            # 获取文件名（不含扩展名）
            base_name = os.path.splitext(os.path.basename(pdf_path))[0]
            output_path = os.path.join(output_dir, f"{base_name}.pptx")
            
            self.log_step("PPT转换", f"开始转换 {base_name}.pdf")
            
            # 使用工厂模式获取转换器
            if hasattr(self, 'converter_factory') and self.converter_factory:
                try:
                    self.log("使用工厂模式获取PDF到PPT转换器")
                    converter = self.converter_factory.get_converter('pdf', 'pptx')
                    if converter:
                        self.log(f"找到转换器: {converter.name}")
                        success = converter.convert(pdf_path, output_path)
                        if success:
                            self.log_success(f"转换完成: {base_name}.pdf -> {base_name}.pptx")
                            self.log(f"输出文件: {output_path}")
                            return output_path
                        else:
                            self.log_error("转换器执行失败", None)
                    else:
                        self.log_error("未找到支持PDF到PPT转换的转换器", None)
                except Exception as e:
                    self.log_error(f"工厂模式转换出错: {e}", e)
            
            raise Exception("没有可用的PDF转PPT转换器")
            
        except Exception as e:
            self.log_error(f"PPT转换失败", e)
            raise e
    
    def convert_to_pptx_via_word(self, pdf_path, output_dir):
        """通过Word中转将PDF转换为PPT"""
        # 获取文件名（不含扩展名）
        base_name = os.path.splitext(os.path.basename(pdf_path))[0]
        output_path = os.path.join(output_dir, f"{base_name}.pptx")
        
        try:
            self.log_step("PDF→Word→PPT转换", f"开始转换 {base_name}.pdf")
            
            # 步骤1: 先将PDF转换为Word
            self.log_step("步骤1", "将PDF转换为Word文档")
            with tempfile.TemporaryDirectory() as temp_dir:
                temp_word_path = os.path.join(temp_dir, f"{base_name}_temp.docx")
                
                # 使用现有的PDF转Word方法
                try:
                    from pdf2docx import parse
                    self.log("开始使用pdf2docx转换为Word")
                    
                    # 获取页面总数用于进度显示
                    try:
                        import fitz
                        doc = fitz.open(pdf_path)
                        total_pages = len(doc)
                        doc.close()
                        self.log(f"PDF总页数: {total_pages}")
                    except:
                        total_pages = 0
                    
                    # 添加更详细的参数配置来处理问题PDF
                    parse(pdf_path, temp_word_path,
                          start=0, end=None,
                          pages=None,
                          password=None,
                          multi_processing=False,
                          cpu_count=1)
                    
                    self.log_success("使用pdf2docx成功转换为Word")
                except Exception as e:
                    error_msg = str(e)
                    self.log_error(f"pdf2docx转换失败", e)
                    
                    # 检查是否是已知的bandwriter错误
                    if "Invalid bandwriter header dimensions" in error_msg or "bandwriter" in error_msg:
                        self.log_fallback("检测到bandwriter错误，尝试逐页转换")
                        if self._try_page_by_page_conversion(pdf_path, temp_word_path):
                            self.log_success("页面范围转换成功")
                        else:
                            self.log_fallback("页面范围转换失败，使用备用方法")
                            self._convert_to_docx_fallback_for_ppt(pdf_path, temp_word_path)
                    else:
                        self.log_fallback("使用备用方法")
                        self._convert_to_docx_fallback_for_ppt(pdf_path, temp_word_path)
                
                # 步骤2: 将Word转换为PPT
                self.log_step("步骤2", "将Word文档转换为PPT")
                
                # 使用工厂模式进行Word到PPT转换
                if hasattr(self, 'converter_factory') and self.converter_factory:
                    try:
                        self.log("使用工厂模式获取Word到PPT转换器")
                        word_converter = self.converter_factory.get_converter('docx', 'pptx')
                        if word_converter:
                            self.log(f"找到Word转换器: {word_converter.name}")
                            word_to_ppt_success = word_converter.convert(temp_word_path, output_path)
                            if word_to_ppt_success:
                                self.log("Word到PPT转换成功")
                                result_path = output_path
                            else:
                                raise Exception("Word到PPT转换失败")
                        else:
                            raise Exception("未找到支持Word到PPT转换的转换器")
                    except Exception as e:
                        self.log_error(f"Word到PPT转换出错: {e}", e)
                        raise e
                else:
                    raise Exception("转换器工厂未初始化")
                
                self.log_success(f"转换完成: {base_name}.pdf -> {base_name}.pptx (通过Word中转)")
                self.log(f"输出文件: {result_path}")
                
        except Exception as e:
            self.log_error(f"PDF→Word→PPT转换失败", e)
            raise e
    
    def _convert_to_docx_fallback_for_ppt(self, pdf_path, output_path):
        """为PPT转换专用的备用PDF转DOCX方法"""
        from docx import Document
        import fitz
        
        self.log_fallback("使用备用方法提取PDF文本")
        
        # 创建一个新的Word文档
        doc = Document()
        
        # 打开PDF文件
        pdf_document = fitz.open(pdf_path)
        total_pages = len(pdf_document)
        self.log(f"PDF总页数: {total_pages}")
        
        # 遍历每一页
        for page_num in range(total_pages):
            self.log_progress(page_num + 1, total_pages, f"提取页面 {page_num + 1} 文本")
            
            page = pdf_document.load_page(page_num)
            text = page.get_text("text")
            
            # 如果页面有文本内容，添加到文档
            if text.strip():
                # 添加页面内容，不添加页码标题（为PPT转换优化）
                paragraphs = text.split('\n')
                for para in paragraphs:
                    if para.strip():
                        doc.add_paragraph(para.strip())
                
                # 页面之间添加分隔
                if page_num < total_pages - 1:
                    doc.add_page_break()
        
        pdf_document.close()
        
        # 保存Word文档
        doc.save(output_path)
        self.log_success(f"备用方法转换完成: {total_pages} 页文本已提取")
    
    def convert_to_image(self, pdf_path, output_dir, image_format, dpi):
        # 获取文件名（不含扩展名）
        base_name = os.path.splitext(os.path.basename(pdf_path))[0]
        
        # 清理文件名中的非法字符，确保可以创建文件夹
        safe_folder_name = self.sanitize_filename(base_name)
        
        # 创建以PDF文件名命名的文件夹
        pdf_output_dir = os.path.join(output_dir, safe_folder_name)
        os.makedirs(pdf_output_dir, exist_ok=True)
        
        self.log_step("图像转换", f"开始转换 {base_name}.pdf 为 {image_format.upper()} (DPI: {dpi})")
        self.log(f"图片将保存到文件夹: {safe_folder_name}")
        
        # 打开PDF文件
        pdf_document = fitz.open(pdf_path)
        total_pages = len(pdf_document)
        self.log(f"PDF总页数: {total_pages}")
        
        # 计算缩放因子
        zoom = dpi / 72  # PDF使用72 DPI作为基准
        self.log(f"缩放因子: {zoom:.2f}")
        
        # 遍历每一页
        for page_num in range(total_pages):
            self.log_progress(page_num + 1, total_pages, f"渲染页面 {page_num + 1}")
            
            try:
                page = pdf_document.load_page(page_num)
                
                # 渲染页面为像素图
                pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom))
                
                # 将像素图转换为PIL图像
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                
                # 保存图像到PDF文件名命名的文件夹中
                output_path = os.path.join(pdf_output_dir, f"{base_name}_page{page_num + 1}.{image_format}")
                
                if image_format.lower() == "jpg":
                    img.save(output_path, "JPEG", quality=95)
                else:  # png
                    img.save(output_path, "PNG")
                    
            except Exception as e:
                self.log_error(f"页面 {page_num + 1} 转换失败", e)
        
        pdf_document.close()
        self.log_success(f"图像转换完成: {total_pages} 页已转换为 {image_format.upper()}")
    
    def setup_convert_mode_ui(self):
        """设置转换模式的UI"""
        pass  # UI已经在setup_ui中设置
    
    def switch_to_convert_mode(self):
        """切换到转换模式"""
        self.current_mode = 'convert'
        
        # 更新按钮样式
        self.pdf_convert_btn.configure(style="Primary.TButton")
        self.pdf_operation_btn.configure(style="Secondary.TButton")
        
        # 更新左侧标题
        self.left_title_label.configure(text="PDF转其他")
        
        # 显示转换功能列表，隐藏操作功能列表
        self.convert_functions_frame.pack(fill=tk.BOTH, expand=True)
        self.operation_functions_frame.pack_forget()
        
        # 显示转换模式UI，隐藏操作模式UI
        self.convert_mode_frame.pack(fill=tk.BOTH, expand=True)
        self.operation_mode_frame.pack_forget()
        
        # 隐藏PDF操作模块的UI
        self.pdf_operations.hide_ui()
    
    def switch_to_operation_mode(self):
        """切换到操作模式"""
        self.current_mode = 'operation'
        
        # 更新按钮样式
        self.pdf_convert_btn.configure(style="Secondary.TButton")
        self.pdf_operation_btn.configure(style="Primary.TButton")
        
        # 更新左侧标题
        self.left_title_label.configure(text="PDF操作")
        
        # 显示操作功能列表，隐藏转换功能列表
        self.operation_functions_frame.pack(fill=tk.BOTH, expand=True)
        self.convert_functions_frame.pack_forget()
        
        # 显示操作模式UI，隐藏转换模式UI
        self.operation_mode_frame.pack(fill=tk.BOTH, expand=True)
        self.convert_mode_frame.pack_forget()
        
        # 显示PDF操作模块的UI
        self.pdf_operations.show_ui(self.operation_mode_frame)
    
    # _convert_to_docx_with_ocr方法已移除（OCR功能已删除）
    
    def on_operation_change(self, *args):
        """当PDF操作功能选择改变时调用"""
        if self.current_mode == 'operation':
            operation = self.selected_operation.get()
            self.pdf_operations.switch_operation(operation)
    
    def _check_dependencies_on_startup(self):
        """应用启动时的依赖检查"""
        try:
            checker = DependencyChecker()
            
            # 执行快速检查（不显示详细报告）
            all_deps_ok = checker.check_all(verbose=False)
            
            if not all_deps_ok:
                # 获取缺失的依赖
                missing = checker.get_missing_dependencies()
                
                # 构建警告消息
                warning_msg = "⚠️ 检测到缺失的依赖项:\n\n"
                
                if missing['python']:
                    warning_msg += "Python包缺失:\n"
                    for pkg in missing['python']:
                        warning_msg += f"  • {pkg}\n"
                    warning_msg += "\n解决方案: pip install -r requirements.txt\n\n"
                
                if missing['system']:
                    warning_msg += "系统依赖缺失:\n"
                    for dep in missing['system']:
                        warning_msg += f"  • {dep}\n"
                    warning_msg += "\n解决方案: 运行 python scripts/setup.py\n\n"
                
                warning_msg += "应用程序将继续运行，但某些功能可能不可用。\n"
                warning_msg += "建议安装缺失的依赖以获得完整功能。"
                
                # 显示警告对话框（延迟到GUI创建后）
                self._show_dependency_warning = warning_msg
            else:
                self._show_dependency_warning = None
                
        except Exception as e:
            print(f"依赖检查时出错: {e}")
            self._show_dependency_warning = None
    
    def _initialize_plugin_system(self):
        """初始化插件系统"""
        try:
            print("正在初始化插件系统...")
            
            # 初始化插件管理器
            self.plugin_manager = get_plugin_manager()
            
            # 加载所有插件
            loaded_count = initialize_plugins()
            print(f"插件系统初始化完成，成功加载 {loaded_count} 个插件")
            
            # 初始化转换器工厂（工厂会自动注册内置转换器和插件转换器）
            self.converter_factory = ConverterFactory.get_instance()
            
            print(f"转换器工厂初始化完成，注册了 {len(self.converter_factory.get_all_converters())} 个转换器")
            
            # 获取支持的格式信息
            self.supported_formats = self.converter_factory.get_supported_formats()
            print(f"支持的输入格式: {self.supported_formats['input']}")
            print(f"支持的输出格式: {self.supported_formats['output']}")
            
        except Exception as e:
            print(f"插件系统初始化失败: {e}")
            # 创建空的工厂作为后备
            self.converter_factory = ConverterFactory()
            self.plugin_manager = None
            self.supported_formats = {'input': ['pdf'], 'output': ['docx', 'pptx']}
    
    def _show_startup_warnings(self):
        """显示启动时的警告信息"""
        if hasattr(self, '_show_dependency_warning') and self._show_dependency_warning:
            # 延迟显示警告，确保GUI已完全加载
            self.root.after(1000, lambda: messagebox.showwarning(
                "依赖检查警告", 
                self._show_dependency_warning
            ))
    
    def run(self):
        # 显示启动警告（如果有）
        self._show_startup_warnings()
        self.root.mainloop()


def main():
    """主函数 - 包含完整的依赖检查和错误处理"""
    try:
        # 执行详细的依赖检查
        print("🚀 启动PDF转换器...")
        print("📋 正在检查依赖项...")
        
        checker = DependencyChecker()
        deps_ok = checker.check_all(verbose=True)
        
        if not deps_ok:
            print("\n⚠️  发现缺失的依赖项，但应用程序将继续运行。")
            print("💡 建议运行 'python scripts/setup.py' 进行自动安装。")
            
            # 询问用户是否继续
            response = input("\n是否继续启动应用程序? (y/n): ")
            if response.lower() != 'y':
                print("应用程序启动已取消。")
                return
        
        print("\n🎉 启动应用程序...")
        app = PDFConverter()
        app.run()
        
    except KeyboardInterrupt:
        print("\n用户中断，应用程序退出。")
    except Exception as e:
        print(f"\n❌ 应用程序启动失败: {e}")
        print("💡 请检查依赖安装或运行 'python scripts/setup.py'")


if __name__ == "__main__":
    main()
